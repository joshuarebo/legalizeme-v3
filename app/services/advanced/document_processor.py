import logging
import asyncio
from typing import List, Dict, Optional, Any, Union
from pathlib import Path
import json
import re
from datetime import datetime
import tempfile
import os

# Optional imports for document processing
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from pypdf2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

try:
    import boto3
    HAS_BOTO3 = True
except ImportError:
    HAS_BOTO3 = False

from app.config import settings
from app.services.ai_service import AIService

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processor with multi-modal capabilities"""
    
    def __init__(self, ai_service: AIService = None):
        self.ai_service = ai_service or AIService()
        self.bedrock_client = None
        
        # Processing configuration
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.supported_formats = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.txt': 'text',
            '.html': 'html',
            '.htm': 'html'
        }
        
        # OCR configuration
        self.ocr_languages = ['eng']  # English by default
        self.ocr_config = '--oem 3 --psm 6'
        
        self._initialized = False
    
    async def initialize(self):
        """Initialize the document processor"""
        if self._initialized:
            return
        
        try:
            logger.info("Initializing Document Processor...")
            
            # Initialize Bedrock client for AI processing
            if HAS_BOTO3:
                try:
                    self.bedrock_client = boto3.client(
                        'bedrock-runtime',
                        region_name=getattr(settings, 'AWS_REGION', 'us-east-1')
                    )
                    logger.info("AWS Bedrock client initialized for document processing")
                except Exception as e:
                    logger.warning(f"Failed to initialize Bedrock client: {e}")
            
            # Check OCR availability
            if HAS_OCR:
                try:
                    # Test OCR installation
                    pytesseract.get_tesseract_version()
                    logger.info("OCR capabilities available")
                except Exception as e:
                    logger.warning(f"OCR not properly configured: {e}")
            
            self._initialized = True
            logger.info("Document Processor initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize Document Processor: {e}")
            raise
    
    async def analyze_pdf(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Analyze PDF document with OCR fallback"""
        if not self._initialized:
            await self.initialize()
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"success": False, "error": "File not found"}
        
        if file_path.stat().st_size > self.max_file_size:
            return {"success": False, "error": "File too large"}
        
        try:
            # Try pdfplumber first (better for text extraction)
            if HAS_PDFPLUMBER:
                result = await self._extract_with_pdfplumber(file_path)
                if result["success"] and result.get("text", "").strip():
                    return result
            
            # Fallback to pypdf2
            if HAS_PYPDF2:
                result = await self._extract_with_pypdf2(file_path)
                if result["success"] and result.get("text", "").strip():
                    return result
            
            # Final fallback to OCR
            if HAS_OCR:
                logger.info(f"Falling back to OCR for {file_path.name}")
                return await self._extract_with_ocr(file_path)
            
            return {"success": False, "error": "No PDF processing libraries available"}
            
        except Exception as e:
            logger.error(f"Error analyzing PDF {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    async def _extract_with_pdfplumber(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using pdfplumber"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                metadata = {
                    "total_pages": len(pdf.pages),
                    "extraction_method": "pdfplumber"
                }
                
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                full_text = "\n".join(text_parts)
                
                return {
                    "success": True,
                    "text": full_text,
                    "metadata": metadata,
                    "word_count": len(full_text.split()),
                    "char_count": len(full_text)
                }
                
        except Exception as e:
            logger.error(f"Error with pdfplumber extraction: {e}")
            return {"success": False, "error": str(e)}
    
    async def _extract_with_pypdf2(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using pypdf2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text_parts = []
                
                metadata = {
                    "total_pages": len(pdf_reader.pages),
                    "extraction_method": "pypdf2"
                }
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                full_text = "\n".join(text_parts)
                
                return {
                    "success": True,
                    "text": full_text,
                    "metadata": metadata,
                    "word_count": len(full_text.split()),
                    "char_count": len(full_text)
                }
                
        except Exception as e:
            logger.error(f"Error with pypdf2 extraction: {e}")
            return {"success": False, "error": str(e)}
    
    async def _extract_with_ocr(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using OCR"""
        try:
            # Convert PDF to images and OCR each page
            import fitz  # PyMuPDF for PDF to image conversion
            
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Save to temporary file
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
                    temp_file.write(img_data)
                    temp_path = temp_file.name
                
                try:
                    # OCR the image
                    text = pytesseract.image_to_string(
                        Image.open(temp_path),
                        lang='+'.join(self.ocr_languages),
                        config=self.ocr_config
                    )
                    if text.strip():
                        text_parts.append(text)
                finally:
                    # Clean up temp file
                    os.unlink(temp_path)
            
            doc.close()
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "total_pages": len(doc),
                    "extraction_method": "ocr",
                    "ocr_languages": self.ocr_languages
                },
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }
            
        except Exception as e:
            logger.error(f"Error with OCR extraction: {e}")
            return {"success": False, "error": str(e)}
    
    async def analyze_docx(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Analyze DOCX document"""
        if not self._initialized:
            await self.initialize()
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"success": False, "error": "File not found"}
        
        try:
            # Try python-docx first
            if HAS_DOCX:
                result = await self._extract_with_docx(file_path)
                if result["success"]:
                    return result
            
            # Fallback to mammoth
            if HAS_MAMMOTH:
                return await self._extract_with_mammoth(file_path)
            
            return {"success": False, "error": "No DOCX processing libraries available"}
            
        except Exception as e:
            logger.error(f"Error analyzing DOCX {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    async def _extract_with_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using python-docx"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": {
                    "extraction_method": "python-docx",
                    "paragraph_count": len(doc.paragraphs)
                },
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }
            
        except Exception as e:
            logger.error(f"Error with python-docx extraction: {e}")
            return {"success": False, "error": str(e)}
    
    async def _extract_with_mammoth(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using mammoth"""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                
            return {
                "success": True,
                "text": result.value,
                "metadata": {
                    "extraction_method": "mammoth",
                    "warnings": result.messages
                },
                "word_count": len(result.value.split()),
                "char_count": len(result.value)
            }
            
        except Exception as e:
            logger.error(f"Error with mammoth extraction: {e}")
            return {"success": False, "error": str(e)}

    async def summarize_text(self, text: str, model: str = "claude-sonnet-4", max_length: int = 500) -> Dict[str, Any]:
        """Summarize text using specified AI model"""
        if not self._initialized:
            await self.initialize()

        if not text or not text.strip():
            return {"success": False, "error": "No text provided"}

        try:
            # Prepare summarization prompt
            prompt = self._create_summarization_prompt(text, max_length)

            # Generate summary using AI service
            if model.startswith("claude") and self.bedrock_client:
                result = await self._summarize_with_bedrock(prompt, model)
            else:
                result = await self._summarize_with_ai_service(prompt)

            if result["success"]:
                # Analyze summary quality
                quality_score = self._assess_summary_quality(text, result["summary"])
                result["quality_score"] = quality_score

            return result

        except Exception as e:
            logger.error(f"Error summarizing text: {e}")
            return {"success": False, "error": str(e)}

    def _create_summarization_prompt(self, text: str, max_length: int) -> str:
        """Create prompt for text summarization"""
        return f"""Please provide a comprehensive summary of the following legal document. Focus on key legal principles, important provisions, and practical implications.

DOCUMENT TEXT:
{text[:4000]}  # Limit input text to avoid token limits

REQUIREMENTS:
1. Maximum length: {max_length} words
2. Focus on legal significance and practical implications
3. Maintain accuracy and legal terminology
4. Include key dates, parties, and legal references if present
5. Structure the summary clearly with main points

SUMMARY:"""

    async def _summarize_with_bedrock(self, prompt: str, model: str) -> Dict[str, Any]:
        """Summarize using AWS Bedrock"""
        try:
            # Map model names to Bedrock model IDs
            model_mapping = {
                "claude-sonnet-4": "us.anthropic.claude-sonnet-4-20250514-v1:0",
                "claude-3.7": "us.anthropic.claude-3-7-sonnet-20250219-v1:0",
                "claude-haiku": "anthropic.claude-3-haiku-20240307-v1:0"
            }

            model_id = model_mapping.get(model, model_mapping["claude-sonnet-4"])

            # Prepare request body
            body = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 1000,
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            }

            # Make request to Bedrock
            response = self.bedrock_client.invoke_model(
                modelId=model_id,
                body=json.dumps(body)
            )

            # Parse response
            result = json.loads(response['body'].read())
            summary = result['content'][0]['text']

            return {
                "success": True,
                "summary": summary,
                "model_used": model,
                "method": "bedrock"
            }

        except Exception as e:
            logger.error(f"Error with Bedrock summarization: {e}")
            return {"success": False, "error": str(e)}

    async def _summarize_with_ai_service(self, prompt: str) -> Dict[str, Any]:
        """Summarize using AI service fallback"""
        try:
            result = await self.ai_service._generate_with_fallback(prompt, "summarization")

            return {
                "success": True,
                "summary": result.get("response", ""),
                "model_used": result.get("model_used", "unknown"),
                "method": "ai_service"
            }

        except Exception as e:
            logger.error(f"Error with AI service summarization: {e}")
            return {"success": False, "error": str(e)}

    def _assess_summary_quality(self, original_text: str, summary: str) -> float:
        """Assess the quality of a summary"""
        try:
            if not summary or not original_text:
                return 0.0

            # Basic quality metrics
            original_words = len(original_text.split())
            summary_words = len(summary.split())

            # Compression ratio (should be reasonable)
            compression_ratio = summary_words / original_words if original_words > 0 else 0
            compression_score = 1.0 if 0.1 <= compression_ratio <= 0.3 else 0.5

            # Check for key legal terms preservation
            legal_terms = ['section', 'article', 'act', 'court', 'judgment', 'law', 'legal', 'provision']
            original_legal_terms = sum(1 for term in legal_terms if term.lower() in original_text.lower())
            summary_legal_terms = sum(1 for term in legal_terms if term.lower() in summary.lower())

            legal_preservation = (
                summary_legal_terms / original_legal_terms
                if original_legal_terms > 0 else 1.0
            )

            # Overall quality score
            quality_score = (compression_score * 0.3 + legal_preservation * 0.7)
            return min(quality_score, 1.0)

        except Exception as e:
            logger.error(f"Error assessing summary quality: {e}")
            return 0.5

    async def generate_embeddings(self, text: str) -> Optional[List[float]]:
        """Generate embeddings for text"""
        if not self._initialized:
            await self.initialize()

        try:
            # Use AI service for embeddings (more reliable)
            return await self.ai_service.generate_embeddings(text)

        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            return None

# Removed Bedrock Titan embedding - using AI service instead for reliability

    async def analyze_document_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure and extract metadata"""
        try:
            analysis = {
                "document_type": "unknown",
                "structure": {},
                "key_sections": [],
                "entities": [],
                "dates": [],
                "legal_references": []
            }

            # Detect document type
            analysis["document_type"] = self._detect_document_type(text)

            # Extract structure based on document type
            if analysis["document_type"] == "legislation":
                analysis["structure"] = self._analyze_legislation_structure(text)
            elif analysis["document_type"] == "judgment":
                analysis["structure"] = self._analyze_judgment_structure(text)
            elif analysis["document_type"] == "constitution":
                analysis["structure"] = self._analyze_constitution_structure(text)

            # Extract common elements
            analysis["dates"] = self._extract_dates(text)
            analysis["legal_references"] = self._extract_legal_references(text)
            analysis["key_sections"] = self._extract_key_sections(text)

            return analysis

        except Exception as e:
            logger.error(f"Error analyzing document structure: {e}")
            return {"error": str(e)}

    def _detect_document_type(self, text: str) -> str:
        """Detect the type of legal document"""
        text_lower = text.lower()

        # Constitution indicators
        if any(term in text_lower for term in ['constitution', 'constitutional', 'bill of rights']):
            return "constitution"

        # Legislation indicators
        elif any(term in text_lower for term in ['act', 'statute', 'law', 'section', 'chapter']):
            return "legislation"

        # Judgment indicators
        elif any(term in text_lower for term in ['judgment', 'ruling', 'court', 'plaintiff', 'defendant']):
            return "judgment"

        # Regulation indicators
        elif any(term in text_lower for term in ['regulation', 'rule', 'order']):
            return "regulation"

        # Gazette indicators
        elif any(term in text_lower for term in ['gazette', 'notice', 'proclamation']):
            return "gazette"

        else:
            return "legal_document"

    def _analyze_legislation_structure(self, text: str) -> Dict[str, Any]:
        """Analyze structure of legislation"""
        structure = {
            "title": "",
            "preamble": "",
            "parts": [],
            "sections": [],
            "schedules": []
        }

        try:
            # Extract title (usually at the beginning)
            title_match = re.search(r'^([A-Z][^.\n]+(?:ACT|LAW|STATUTE))', text, re.MULTILINE | re.IGNORECASE)
            if title_match:
                structure["title"] = title_match.group(1).strip()

            # Extract sections
            section_pattern = r'(?:^|\n)\s*(\d+)\.\s+([^\n]+)'
            sections = re.findall(section_pattern, text, re.MULTILINE)
            structure["sections"] = [
                {"number": num, "title": title.strip()}
                for num, title in sections[:50]  # Limit to first 50 sections
            ]

            # Extract parts
            part_pattern = r'(?:^|\n)\s*PART\s+([IVX\d]+)\s*[-–—]\s*([^\n]+)'
            parts = re.findall(part_pattern, text, re.MULTILINE | re.IGNORECASE)
            structure["parts"] = [
                {"number": num, "title": title.strip()}
                for num, title in parts
            ]

        except Exception as e:
            logger.error(f"Error analyzing legislation structure: {e}")

        return structure

    def _analyze_judgment_structure(self, text: str) -> Dict[str, Any]:
        """Analyze structure of court judgment"""
        structure = {
            "case_title": "",
            "court": "",
            "judges": [],
            "parties": {},
            "case_number": "",
            "date": ""
        }

        try:
            # Extract case title
            case_pattern = r'([A-Z][A-Z\s&]+)\s+(?:vs?\.?|V\.?)\s+([A-Z][A-Z\s&]+)'
            case_match = re.search(case_pattern, text)
            if case_match:
                structure["case_title"] = f"{case_match.group(1).strip()} v. {case_match.group(2).strip()}"
                structure["parties"] = {
                    "plaintiff": case_match.group(1).strip(),
                    "defendant": case_match.group(2).strip()
                }

            # Extract court
            court_pattern = r'(?:IN THE|BEFORE THE)\s+(.*?(?:COURT|TRIBUNAL))'
            court_match = re.search(court_pattern, text, re.IGNORECASE)
            if court_match:
                structure["court"] = court_match.group(1).strip()

            # Extract case number
            case_num_pattern = r'(?:Case No\.?|Cause No\.?|Petition No\.?)\s*:?\s*([A-Z0-9\/\-\s]+)'
            case_num_match = re.search(case_num_pattern, text, re.IGNORECASE)
            if case_num_match:
                structure["case_number"] = case_num_match.group(1).strip()

        except Exception as e:
            logger.error(f"Error analyzing judgment structure: {e}")

        return structure

    def _analyze_constitution_structure(self, text: str) -> Dict[str, Any]:
        """Analyze structure of constitution"""
        structure = {
            "chapters": [],
            "articles": [],
            "bill_of_rights": False
        }

        try:
            # Extract chapters
            chapter_pattern = r'(?:^|\n)\s*CHAPTER\s+([IVX\d]+)\s*[-–—]\s*([^\n]+)'
            chapters = re.findall(chapter_pattern, text, re.MULTILINE | re.IGNORECASE)
            structure["chapters"] = [
                {"number": num, "title": title.strip()}
                for num, title in chapters
            ]

            # Extract articles
            article_pattern = r'(?:^|\n)\s*Article\s+(\d+)\.\s+([^\n]+)'
            articles = re.findall(article_pattern, text, re.MULTILINE | re.IGNORECASE)
            structure["articles"] = [
                {"number": num, "title": title.strip()}
                for num, title in articles[:100]  # Limit to first 100 articles
            ]

            # Check for Bill of Rights
            structure["bill_of_rights"] = "bill of rights" in text.lower()

        except Exception as e:
            logger.error(f"Error analyzing constitution structure: {e}")

        return structure

    def _extract_dates(self, text: str) -> List[str]:
        """Extract dates from text"""
        date_patterns = [
            r'\b\d{1,2}(?:st|nd|rd|th)?\s+\w+\s+\d{4}\b',  # 1st January 2020
            r'\b\w+\s+\d{1,2},?\s+\d{4}\b',  # January 1, 2020
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b',  # 01/01/2020
            r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'   # 2020/01/01
        ]

        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)

        return list(set(dates))[:10]  # Return unique dates, limit to 10

    def _extract_legal_references(self, text: str) -> List[str]:
        """Extract legal references from text"""
        reference_patterns = [
            r'\b[A-Z][a-z]+\s+Act\s+(?:No\.?\s*)?\d+\s+of\s+\d{4}\b',  # Employment Act No. 11 of 2007
            r'\bArticle\s+\d+\b',  # Article 25
            r'\bSection\s+\d+\b',  # Section 10
            r'\bChapter\s+\d+\b',  # Chapter 5
            r'\b[A-Z][a-z]+\s+v\.?\s+[A-Z][a-z]+\b'  # Case citations
        ]

        references = []
        for pattern in reference_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            references.extend(matches)

        return list(set(references))[:20]  # Return unique references, limit to 20

    def _extract_key_sections(self, text: str) -> List[Dict[str, str]]:
        """Extract key sections from document"""
        sections = []

        # Look for numbered sections
        section_pattern = r'(?:^|\n)\s*(\d+)\.\s+([^\n]+(?:\n(?!\s*\d+\.)[^\n]*)*)'
        matches = re.findall(section_pattern, text, re.MULTILINE)

        for number, content in matches[:10]:  # Limit to first 10 sections
            sections.append({
                "number": number,
                "title": content.split('\n')[0].strip(),
                "content": content.strip()[:200] + "..." if len(content) > 200 else content.strip()
            })

        return sections

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())

    def get_capabilities(self) -> Dict[str, bool]:
        """Get processor capabilities"""
        return {
            "pdf_extraction": HAS_PDFPLUMBER or HAS_PYPDF2,
            "docx_extraction": HAS_DOCX or HAS_MAMMOTH,
            "ocr_processing": HAS_OCR,
            "ai_summarization": True,
            "bedrock_integration": HAS_BOTO3 and self.bedrock_client is not None,
            "structure_analysis": True,
            "embedding_generation": True
        }
