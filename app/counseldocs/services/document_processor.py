"""
Document Processor Service for CounselDocs
Handles multi-format document processing with OCR and content extraction.
Production-ready for Kenyan legal landscape with large document support.
"""

import asyncio
import logging
import mimetypes
import os
import tempfile
from typing import Dict, Any, Optional, Tuple, List
from datetime import datetime
import uuid

import boto3
from botocore.exceptions import ClientError
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

from app.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Production-ready document processor for Kenyan legal documents.
    Supports large files up to 100MB (typical for multi-page contracts).
    """
    
    # Production file size limits for Kenyan legal landscape
    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB for large contracts/agreements
    SUPPORTED_FORMATS = {
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'doc': ['application/msword'],
        'txt': ['text/plain'],
        'png': ['image/png'],
        'jpg': ['image/jpeg'],
        'jpeg': ['image/jpeg'],
        'tiff': ['image/tiff'],
        'bmp': ['image/bmp']
    }
    
    def __init__(self):
        """Initialize document processor with AWS services"""
        self.s3_client = boto3.client(
            's3',
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_REGION
        )
        
        self.textract_client = boto3.client(
            'textract',
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_REGION
        )
        
        # S3 bucket for document storage
        self.s3_bucket = "counseldocs-documents"

        # Try to ensure bucket exists, but don't fail startup if it doesn't work
        try:
            self._ensure_s3_bucket()
        except Exception as e:
            logger.warning(f"S3 bucket initialization failed: {e}. Will try again during first upload.")
    
    def _ensure_s3_bucket(self):
        """Ensure S3 bucket exists for document storage"""
        try:
            self.s3_client.head_bucket(Bucket=self.s3_bucket)
        except ClientError as e:
            if e.response['Error']['Code'] == '404':
                try:
                    # For us-east-1, don't specify LocationConstraint
                    if settings.AWS_REGION == 'us-east-1':
                        self.s3_client.create_bucket(Bucket=self.s3_bucket)
                    else:
                        self.s3_client.create_bucket(
                            Bucket=self.s3_bucket,
                            CreateBucketConfiguration={'LocationConstraint': settings.AWS_REGION}
                        )
                    logger.info(f"Created S3 bucket: {self.s3_bucket}")
                except ClientError as create_error:
                    logger.error(f"Failed to create S3 bucket: {create_error}")
                    raise
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str, 
        content_type: str,
        user_id: str
    ) -> Dict[str, Any]:
        """
        Process uploaded document with content extraction.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type
            user_id: Azure user ID
            
        Returns:
            Dict with processing results
        """
        start_time = datetime.utcnow()
        
        try:
            # Validate file
            validation_result = self._validate_file(file_content, filename, content_type)
            if not validation_result['valid']:
                return {
                    'success': False,
                    'error': validation_result['error'],
                    'error_type': 'validation'
                }
            
            file_type = validation_result['file_type']
            file_size = len(file_content)
            
            # Upload to S3
            s3_path = await self._upload_to_s3(file_content, filename, user_id, file_type)
            
            # Extract content based on file type
            extraction_result = await self._extract_content(file_content, file_type, filename)
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            return {
                'success': True,
                'file_info': {
                    'original_filename': filename,
                    'file_type': file_type,
                    'file_size_bytes': file_size,
                    's3_file_path': s3_path,
                    'content_type': content_type
                },
                'extraction_result': extraction_result,
                'processing_time_seconds': processing_time,
                'processed_at': start_time.isoformat()
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            return {
                'success': False,
                'error': f"Processing failed: {str(e)}",
                'error_type': 'processing'
            }
    
    def _validate_file(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Validate uploaded file"""
        
        # Check file size
        if len(file_content) > self.MAX_FILE_SIZE:
            return {
                'valid': False,
                'error': f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)}MB"
            }
        
        # Check file type
        file_extension = os.path.splitext(filename)[1].lower().lstrip('.')
        
        if file_extension not in self.SUPPORTED_FORMATS:
            return {
                'valid': False,
                'error': f"Unsupported file type: {file_extension}. Supported: {list(self.SUPPORTED_FORMATS.keys())}"
            }
        
        # Validate MIME type
        expected_mimes = self.SUPPORTED_FORMATS[file_extension]
        if content_type not in expected_mimes:
            # Try to guess from content
            guessed_type, _ = mimetypes.guess_type(filename)
            if guessed_type not in expected_mimes:
                logger.warning(f"MIME type mismatch: {content_type} vs expected {expected_mimes}")
        
        return {
            'valid': True,
            'file_type': file_extension
        }
    
    async def _upload_to_s3(self, file_content: bytes, filename: str, user_id: str, file_type: str) -> str:
        """Upload file to S3 and return path"""
        
        # Generate unique S3 key
        file_uuid = str(uuid.uuid4())
        s3_key = f"users/{user_id}/documents/{file_uuid}_{filename}"
        
        try:
            self.s3_client.put_object(
                Bucket=self.s3_bucket,
                Key=s3_key,
                Body=file_content,
                ContentType=mimetypes.guess_type(filename)[0] or 'application/octet-stream',
                Metadata={
                    'user_id': user_id,
                    'original_filename': filename,
                    'file_type': file_type,
                    'upload_timestamp': datetime.utcnow().isoformat()
                }
            )
            
            logger.info(f"Uploaded file to S3: {s3_key}")
            return s3_key
            
        except ClientError as e:
            logger.error(f"S3 upload failed: {e}")
            raise Exception(f"Failed to upload file: {str(e)}")
    
    async def _extract_content(self, file_content: bytes, file_type: str, filename: str) -> Dict[str, Any]:
        """Extract text content from document"""
        
        try:
            if file_type == 'pdf':
                return await self._extract_pdf_content(file_content)
            elif file_type in ['docx', 'doc']:
                return await self._extract_docx_content(file_content)
            elif file_type == 'txt':
                return await self._extract_text_content(file_content)
            elif file_type in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                return await self._extract_image_content(file_content, file_type)
            else:
                return {
                    'success': False,
                    'error': f"Unsupported file type for extraction: {file_type}"
                }
                
        except Exception as e:
            logger.error(f"Content extraction failed for {filename}: {str(e)}")
            return {
                'success': False,
                'error': f"Content extraction failed: {str(e)}"
            }
    
    async def _extract_pdf_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from PDF using multiple methods"""
        
        extracted_text = ""
        page_count = 0
        extraction_method = "unknown"
        
        try:
            # Method 1: Try pdfplumber (better for complex layouts)
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                try:
                    with pdfplumber.open(temp_file.name) as pdf:
                        page_count = len(pdf.pages)
                        text_parts = []
                        
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                        
                        if text_parts:
                            extracted_text = "\n\n".join(text_parts)
                            extraction_method = "pdfplumber"
                        
                except Exception as e:
                    logger.warning(f"pdfplumber extraction failed: {e}")
                
                # Method 2: Fallback to PyPDF2
                if not extracted_text:
                    try:
                        with open(temp_file.name, 'rb') as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            page_count = len(pdf_reader.pages)
                            text_parts = []
                            
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text_parts.append(page_text)
                            
                            if text_parts:
                                extracted_text = "\n\n".join(text_parts)
                                extraction_method = "PyPDF2"
                                
                    except Exception as e:
                        logger.warning(f"PyPDF2 extraction failed: {e}")
                
                # Method 3: AWS Textract for scanned PDFs
                if not extracted_text or len(extracted_text.strip()) < 100:
                    try:
                        textract_result = await self._extract_with_textract(file_content)
                        if textract_result['success']:
                            extracted_text = textract_result['text']
                            extraction_method = "AWS Textract"
                    except Exception as e:
                        logger.warning(f"Textract extraction failed: {e}")
                
                # Cleanup
                os.unlink(temp_file.name)
        
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return {
                'success': False,
                'error': f"PDF extraction failed: {str(e)}"
            }
        
        return {
            'success': True,
            'extracted_text': extracted_text,
            'page_count': page_count,
            'extraction_method': extraction_method,
            'character_count': len(extracted_text),
            'word_count': len(extracted_text.split()) if extracted_text else 0
        }

    async def _extract_docx_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from DOCX files"""

        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()

                doc = DocxDocument(temp_file.name)
                text_parts = []

                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))

                extracted_text = "\n\n".join(text_parts)

                # Cleanup
                os.unlink(temp_file.name)

                return {
                    'success': True,
                    'extracted_text': extracted_text,
                    'extraction_method': 'python-docx',
                    'character_count': len(extracted_text),
                    'word_count': len(extracted_text.split()) if extracted_text else 0,
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                }

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {
                'success': False,
                'error': f"DOCX extraction failed: {str(e)}"
            }

    async def _extract_text_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from plain text files"""

        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            extracted_text = None

            for encoding in encodings:
                try:
                    extracted_text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if extracted_text is None:
                return {
                    'success': False,
                    'error': 'Unable to decode text file with supported encodings'
                }

            return {
                'success': True,
                'extracted_text': extracted_text,
                'extraction_method': 'text_decode',
                'character_count': len(extracted_text),
                'word_count': len(extracted_text.split()),
                'line_count': len(extracted_text.splitlines())
            }

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return {
                'success': False,
                'error': f"Text extraction failed: {str(e)}"
            }

    async def _extract_image_content(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """Extract text from images using OCR"""

        try:
            # First try AWS Textract (more accurate)
            textract_result = await self._extract_with_textract(file_content)
            if textract_result['success'] and textract_result['text'].strip():
                return {
                    'success': True,
                    'extracted_text': textract_result['text'],
                    'extraction_method': 'AWS Textract',
                    'character_count': len(textract_result['text']),
                    'word_count': len(textract_result['text'].split()),
                    'confidence_score': textract_result.get('confidence', 0.0)
                }

            # Fallback to pytesseract
            with tempfile.NamedTemporaryFile(suffix=f'.{file_type}', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()

                image = Image.open(temp_file.name)
                extracted_text = pytesseract.image_to_string(image)

                # Cleanup
                os.unlink(temp_file.name)

                return {
                    'success': True,
                    'extracted_text': extracted_text,
                    'extraction_method': 'pytesseract',
                    'character_count': len(extracted_text),
                    'word_count': len(extracted_text.split()) if extracted_text else 0,
                    'image_size': image.size
                }

        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            return {
                'success': False,
                'error': f"Image OCR failed: {str(e)}"
            }

    async def _extract_with_textract(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text using AWS Textract"""

        try:
            response = self.textract_client.detect_document_text(
                Document={'Bytes': file_content}
            )

            text_parts = []
            confidence_scores = []

            for block in response['Blocks']:
                if block['BlockType'] == 'LINE':
                    text_parts.append(block['Text'])
                    confidence_scores.append(block.get('Confidence', 0.0))

            extracted_text = "\n".join(text_parts)
            avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0

            return {
                'success': True,
                'text': extracted_text,
                'confidence': avg_confidence / 100.0,  # Convert to 0-1 scale
                'block_count': len([b for b in response['Blocks'] if b['BlockType'] == 'LINE'])
            }

        except ClientError as e:
            logger.error(f"Textract failed: {e}")
            return {
                'success': False,
                'error': f"Textract failed: {str(e)}"
            }

    async def get_document_from_s3(self, s3_path: str) -> Optional[bytes]:
        """Retrieve document from S3"""

        try:
            response = self.s3_client.get_object(Bucket=self.s3_bucket, Key=s3_path)
            return response['Body'].read()
        except ClientError as e:
            logger.error(f"Failed to retrieve document from S3: {e}")
            return None

    async def delete_document_from_s3(self, s3_path: str) -> bool:
        """Delete document from S3"""

        try:
            self.s3_client.delete_object(Bucket=self.s3_bucket, Key=s3_path)
            logger.info(f"Deleted document from S3: {s3_path}")
            return True
        except ClientError as e:
            logger.error(f"Failed to delete document from S3: {e}")
            return False

# Global instance
document_processor = DocumentProcessor()
