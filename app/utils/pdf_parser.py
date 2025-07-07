import logging
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document
import mammoth
import io
from pathlib import Path

logger = logging.getLogger(__name__)

class PDFParser:
    """Utility class for parsing PDF and document files"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """Extract text from a file based on its extension"""
        try:
            file_path = Path(file_path)
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif extension == '.docx':
                return self.extract_text_from_docx(file_path)
            elif extension == '.doc':
                return self.extract_text_from_doc(file_path)
            elif extension == '.txt':
                return self.extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file format: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def extract_text_from_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n"
                
                return text.strip()
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return None
    
    def extract_text_from_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return None
    
    def extract_text_from_doc(self, file_path: Path) -> Optional[str]:
        """Extract text from DOC file using mammoth"""
        try:
            with open(file_path, 'rb') as file:
                result = mammoth.extract_raw_text(file)
                return result.value
                
        except Exception as e:
            logger.error(f"Error extracting text from DOC {file_path}: {e}")
            return None
    
    def extract_text_from_txt(self, file_path: Path) -> Optional[str]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error extracting text from TXT {file_path}: {e}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            return None
    
    def extract_text_from_bytes(self, file_bytes: bytes) -> Optional[str]:
        """Extract text from PDF bytes"""
        try:
            file_stream = io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(file_stream)
            text = ""
            
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF bytes: {e}")
            return None
    
    def extract_text_from_docx_bytes(self, file_bytes: bytes) -> Optional[str]:
        """Extract text from DOCX bytes"""
        try:
            file_stream = io.BytesIO(file_bytes)
            doc = Document(file_stream)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX bytes: {e}")
            return None
    
    def extract_text_from_doc_bytes(self, file_bytes: bytes) -> Optional[str]:
        """Extract text from DOC bytes using mammoth"""
        try:
            file_stream = io.BytesIO(file_bytes)
            result = mammoth.extract_raw_text(file_stream)
            return result.value
            
        except Exception as e:
            logger.error(f"Error extracting text from DOC bytes: {e}")
            return None
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        try:
            file_path = Path(file_path)
            extension = file_path.suffix.lower()
            
            metadata = {
                'filename': file_path.name,
                'extension': extension,
                'size': file_path.stat().st_size,
                'created_at': file_path.stat().st_ctime,
                'modified_at': file_path.stat().st_mtime
            }
            
            if extension == '.pdf':
                metadata.update(self._get_pdf_metadata(file_path))
            elif extension == '.docx':
                metadata.update(self._get_docx_metadata(file_path))
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            return {}
    
    def _get_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Get PDF specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'page_count': len(reader.pages),
                    'title': reader.metadata.get('/Title', '') if reader.metadata else '',
                    'author': reader.metadata.get('/Author', '') if reader.metadata else '',
                    'subject': reader.metadata.get('/Subject', '') if reader.metadata else '',
                    'creator': reader.metadata.get('/Creator', '') if reader.metadata else '',
                    'producer': reader.metadata.get('/Producer', '') if reader.metadata else '',
                    'creation_date': reader.metadata.get('/CreationDate', '') if reader.metadata else '',
                    'modification_date': reader.metadata.get('/ModDate', '') if reader.metadata else ''
                }
                
                return metadata
                
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return {}
    
    def _get_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Get DOCX specific metadata"""
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {}
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_formats
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 10) -> bool:
        """Validate file size"""
        try:
            file_size = Path(file_path).stat().st_size
            max_size_bytes = max_size_mb * 1024 * 1024
            return file_size <= max_size_bytes
        except Exception as e:
            logger.error(f"Error validating file size: {e}")
            return False
    
    def validate_file_content(self, file_path: str) -> bool:
        """Validate that file contains extractable text"""
        try:
            text = self.extract_text_from_file(file_path)
            return text is not None and len(text.strip()) > 0
        except Exception as e:
            logger.error(f"Error validating file content: {e}")
            return False
